"""使用 python-docx 生成 Word 教案"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime


def generate_docx(intent: dict, rag_chunks: list[str], output_path: str) -> str:
    topic = intent.get("topic", "未知主题")
    audience = intent.get("audience", "学生")
    key_points = intent.get("key_points", ["核心概念", "基本原理", "实际应用"])
    duration = intent.get("duration", "45分钟")
    style = intent.get("style", "简洁学术")
    today = datetime.date.today().strftime("%Y年%m月%d日")

    doc = Document()

    # 标题
    title = doc.add_heading(f"《{topic}》教学设计方案", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表格
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    data = [
        ("课程主题", topic),
        ("目标受众", audience),
        ("计划课时", duration),
        ("课件风格", style),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()

    # 一、教学目标
    doc.add_heading("一、教学目标", level=1)
    goals = [
        f"知识目标：理解{topic}的基本概念、原理和应用场景",
        f"能力目标：能够运用{topic}相关知识分析和解决实际问题",
        "情感目标：培养学生对本课程的学习兴趣与探究精神",
    ]
    for g in goals:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(g)

    # 二、重点与难点
    doc.add_heading("二、重点与难点", level=1)
    doc.add_paragraph("【教学重点】")
    for kp in key_points[:2]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(kp)
    doc.add_paragraph("【教学难点】")
    for kp in key_points[2:] or key_points[-1:]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(kp)

    # 三、教学步骤
    doc.add_heading("三、教学步骤", level=1)
    steps = [
        ("导入（5分钟）", f"通过提问或案例引入{topic}的概念，激发学生兴趣"),
        ("新课讲授（25分钟）", "逐一讲解各知识点，结合板书和课件演示"),
        ("互动讨论（10分钟）", "学生分组讨论，教师巡回指导"),
        ("课堂小结（5分钟）", "梳理本节课核心内容，强调重难点"),
    ]
    for step_title, step_content in steps:
        doc.add_heading(step_title, level=2)
        doc.add_paragraph(step_content)

    # 四、知识点详解
    doc.add_heading("四、知识点详解", level=1)
    for i, kp in enumerate(key_points):
        doc.add_heading(f"{i+1}. {kp}", level=2)
        doc.add_paragraph(f"{kp}是本课程的重要组成部分。")
        if i < len(rag_chunks) and rag_chunks[i]:
            doc.add_paragraph("【参考资料摘录】")
            p = doc.add_paragraph(rag_chunks[i][:300])
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 五、板书设计
    doc.add_heading("五、板书设计", level=1)
    doc.add_paragraph(f"课题：{topic}")
    for i, kp in enumerate(key_points):
        doc.add_paragraph(f"  {i+1}. {kp}")

    # 六、作业布置
    doc.add_heading("六、作业布置", level=1)
    assignments = [
        f"复习本节课{topic}的核心概念，整理笔记",
        f"完成课后练习题第1-3题",
        f"预习下一节内容，思考{topic}在实际中的应用案例",
    ]
    for a in assignments:
        p = doc.add_paragraph(style="List Number")
        p.add_run(a)

    # 页脚
    doc.add_paragraph()
    footer_p = doc.add_paragraph(f"本教案由 TeachMind AI 智能备课系统生成 · {today}")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    return output_path
